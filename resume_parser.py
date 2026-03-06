import os
import re
import json
import psycopg2
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
import win32com.client

# -------- CONFIGURATION --------
TARGET_FOLDER = r"C:\Users\hkyar\OneDrive\Desktop\Resume parser"
DATABASE_CONFIG = {
    "host": "localhost",
    "database": "resume_parser_db",
    "user": "postgres",
    "password": "2004",
    "port": "5432"
}

# -------- SPECIALIZED READERS --------

def read_pdf(path):
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text: text += page_text + "\n"
    except Exception as e: print(f"   PDF Error: {e}")
    return text

def read_docx(path):
    try:
        doc = Document(path)
        header_text = ""
        for section in doc.sections:
            for para in section.header.paragraphs:
                header_text += para.text + "\n"
        body_text = "\n".join([para.text for para in doc.paragraphs])
        return header_text + body_text
    except Exception as e: print(f"   DOCX Error: {e}"); return ""

def read_doc_legacy(path):
    word = None
    combined_text = ""
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(path, ReadOnly=True)
        for section in doc.Sections:
            for header in section.Headers:
                combined_text += header.Range.Text + "\n"
        combined_text += doc.Content.Text
        doc.Close(False)
    except Exception as e: print(f"   DOC Error: {e}")
    finally:
        if word: word.Quit()
    return combined_text

def read_html(path):
    try:
        with open(path, "r", encoding="utf-8") as file:
            soup = BeautifulSoup(file.read(), "html.parser")
            return soup.get_text(separator="\n")
    except Exception as e: print(f"   HTML Error: {e}"); return ""

# -------- MAIN BATCH LOOP --------

print(f" Scanning folder: {TARGET_FOLDER}\n")

# Connect to DB once for the whole batch
try:
    conn = psycopg2.connect(**DATABASE_CONFIG)
    cursor = conn.cursor()
except Exception as e:
    print(f" Could not connect to Database: {e}")
    exit()

# Loop through every file in the folder
for filename in os.listdir(TARGET_FOLDER):
    file_path = os.path.join(TARGET_FOLDER, filename)
    ext = os.path.splitext(filename)[1].lower()
    
    # Skip folders or non-resume files
    if ext not in [".pdf", ".doc", ".docx", ".html"]:
        continue

    print(f" Processing: {filename}...")
    
    raw_text = ""
    if ext == ".pdf": raw_text = read_pdf(file_path)
    elif ext == ".docx": raw_text = read_docx(file_path)
    elif ext == ".doc": raw_text = read_doc_legacy(file_path)
    elif ext == ".html": raw_text = read_html(file_path)

    # -------- EXTRACTION LOGIC --------
    normalized_text = raw_text.replace('\r', '\n').replace('\u0007', '').replace('\x0b', '\n')
    clean_text = re.sub(r'\s+', ' ', normalized_text)
    lines = [line.strip() for line in normalized_text.split('\n') if line.strip()]

    # 1. Name Extraction
    name = "Not found"
    headers_to_skip = ["key hiring assets", "resume", "cv", "expert in", "developed", "worked on", "education", "profile"]
    for line in lines:
        if not line or len(line) < 3 or len(line) > 50: continue
        if any(h in line.lower() for h in headers_to_skip): continue
        if any(char.isdigit() for char in line): continue
        name = line
        break

    # 2. Education
    education = "Not found"
    edu_match = re.search(r"M\.?C\.?A|B\.?Tech|MBA|B\.E|M\.Tech|BSc", clean_text, re.IGNORECASE)
    if edu_match: education = edu_match.group()

    # 3. Job Title
    job_title = "Not found"
    role_match = re.search(r"Role:\s*([^\n\r]+)", normalized_text)
    if role_match: job_title = role_match.group(1).strip()

    # -------- DATABASE INSERT --------
    try:
        query = "INSERT INTO resumes (name, job_title, education) VALUES (%s, %s, %s)"
        cursor.execute(query, (name, job_title, education))
        conn.commit()
        print(f"    Saved: {name} ({ext})")
    except Exception as e:
        conn.rollback()
        print(f"    DB Error for {filename}: {e}")

# Cleanup
cursor.close()
conn.close()
print("\n All resumes processed!")